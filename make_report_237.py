"""
2.3.7 AI Agent 구축 및 활용 방안 Word 문서 생성
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 설정 (A4) ─────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 스타일 헬퍼 ──────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0x52, 0xFF)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def img_placeholder(label, width_cm=14, height_cm=7):
    """회색 박스 + 텍스트로 사진 삽입 위치 표시"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Cm(width_cm)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F0FE')
    tcPr.append(shd)
    borders_el = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '12')
        b.set(qn('w:color'), '1F6FEB')
        borders_el.append(b)
    tcPr.append(borders_el)

    # 높이
    trPr = tbl.rows[0]._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(height_cm * 5)

    icon_run = cp.add_run('📷  ')
    icon_run.font.size = Pt(14)
    lbl_run = cp.add_run(f'[ 사진 삽입 ]  {label}')
    lbl_run.bold = True
    lbl_run.font.size = Pt(11)
    lbl_run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6E, 0x76, 0x81)
    run.italic = True

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'DDDDDD')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def colored_box(text, bg='EFF6FF', border='1F6FEB'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg)
    tcPr.append(shd)
    borders_el = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '8')
        b.set(qn('w:color'), border)
        borders_el.append(b)
    tcPr.append(borders_el)
    for line in text.split('\n'):
        cp = cell.add_paragraph()
        cp.paragraph_format.left_indent  = Cm(0.3)
        cp.paragraph_format.right_indent = Cm(0.3)
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(2)
        r = cp.add_run(line)
        r.font.size = Pt(10)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ════════════════════════════════════════════════════════════════════════════
#  2.3.7  향후 마케팅 개선을 위한 AI Agent 구축 및 활용 방안
# ════════════════════════════════════════════════════════════════════════════

h1("2.3.7  향후 마케팅 개선을 위한 AI Agent 구축 및 활용 방안")

divider()

# ── 1. 배경 ──────────────────────────────────────────────────────────────────
h2("1.  Agent 개발 배경 — 왜 대고객 채널이 중심인가")

body(
    "앞선 채널 분석(2.3.3~2.3.5)에서 확인했듯, 은행·증권사 채널은 기관 순매수에 직접적인 영향을 미치지만 "
    "운용사가 직접 마케팅 내용을 통제하거나 신속하게 바꾸기 어렵다는 구조적 한계가 있다. "
    "반면 대고객(개인) 채널은 삼성자산운용이 직접 콘텐츠·이벤트·프로모션을 기획·집행할 수 있으며, "
    "개인 투자자의 ETF 순매수 비중이 꾸준히 증가하는 추세 속에서 마케팅 효과가 가장 직접적으로 나타나는 채널이다."
)
body(
    "그러나 대고객 마케팅 모니터링은 수집해야 할 채널이 분산되어 있고(유튜브·이벤트 페이지·뉴스·경쟁사 SNS 등), "
    "이를 매주 수동으로 확인하고 경쟁사와 비교하는 작업은 현실적으로 불가능에 가깝다. "
    "이러한 문제를 해결하기 위해 본 팀은 AI Agent를 직접 설계·개발하였다."
)

divider()

# ── 2. Agent 전체 구조 ────────────────────────────────────────────────────────
h2("2.  AI Agent 전체 구조 — 수집부터 리포트까지 4단계")

body("아래 구조도는 본 팀이 개발한 AI Agent의 전체 작동 흐름이다.")

img_placeholder("[ AI Agent 전체 구조도 — 구조도 이미지 삽입 ]", width_cm=15, height_cm=8)
caption("그림 1.  KODEX ETF 마케팅 모니터링 AI Agent 구조도 (직접 개발)")

body(
    "Agent는 ① 데이터 수집 → ② Claude LLM 분석 → ③ DiD 효과 측정 → ④ 주간 리포트 생성의 4단계로 작동하며, "
    "매주 금요일 장 마감 후 자동 수집이 실행된다."
)

divider()

# ── 3. 단계별 상세 ────────────────────────────────────────────────────────────
h2("3.  단계별 상세 설명")

h3("① 1단계 — 채널별 자동 수집 (4개 세션)")

body(
    "Agent는 매주 4개 채널 그룹에서 데이터를 자동 수집한다. "
    "각 채널은 독립적인 수집기(Collector)로 구성되어 있으며, 랜딩 페이지의 버튼 하나로 전체 수집이 동시 실행된다."
)

tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_data = [('채널', '수집 소스', '수집 항목')]
rows_data = [
    ('📈 증권사 채널', '삼성/미래에셋/키움/토스증권\n한투/신한/KB 등 12개', 'YouTube 영상·이벤트 페이지\n블로그·뉴스 기사'),
    ('🏦 은행 채널',  'KB·신한·하나·우리·NH 9개', 'YouTube 영상·블로그·뉴스\nETF 관련 콘텐츠'),
    ('🎯 대고객 채널', '삼성자산운용 직접 채널\n(KODEX 공식)', 'KODEX YouTube·이벤트 페이지\n네이버 뉴스'),
    ('🏢 경쟁사 채널', 'TIGER·ACE·RISE\nHANARO·SOL', 'YouTube·이벤트 배너\n마케팅 콘텐츠'),
]

# 헤더
for i, h in enumerate(hdr_data[0]):
    cell = tbl.cell(0, i)
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(10)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F6FEB')
    tcPr.append(shd)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for ri, row_data in enumerate(rows_data):
    for ci, txt in enumerate(row_data):
        cell = tbl.cell(ri + 1, ci)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

img_placeholder("[ 랜딩 페이지 — 전체 채널 수집 버튼 화면 ]", width_cm=14, height_cm=6)
caption("그림 2.  랜딩 페이지: 버튼 하나로 4개 세션 전체 수집 실행")

h3("② 2단계 — Claude LLM 마케팅 분석")

body(
    "수집된 텍스트·이미지 데이터를 Claude AI(claude-haiku-4-5 모델)가 분석한다. "
    "단순 키워드 매칭이 아닌 LLM의 문맥 이해를 통해 마케팅 활동 여부를 판단하고, "
    "이벤트 제목·기간·대상 ETF·혜택 조건을 구조화된 JSON으로 추출한다."
)

bullet("증권사 세션: 이벤트·유튜브에서 KODEX ETF 대상 마케팅 감지 → 대상 ETF 종목코드 특정")
bullet("은행 세션: 채널 콘텐츠에서 ETF 마케팅 활동 감지 여부 판단")
bullet("대고객 세션: 배너 이미지(Vision API) + 텍스트 동시 분석 → 이벤트 기간·혜택 추출")
bullet("경쟁사 세션: TIGER·ACE·SOL 등 경쟁사 ETF 마케팅 활동 분류 및 카드보드 생성")

body("분석 결과는 채널 아카이브에 자동 저장되어, 다음 접속 시 버튼 없이 바로 결과가 표시된다.")

img_placeholder("[ 증권사 세션 — 이벤트 카드보드 화면 ]", width_cm=14, height_cm=6)
caption("그림 3.  증권사 채널 LLM 분석 결과: 감지된 ETF 마케팅 이벤트 카드보드")

img_placeholder("[ 경쟁사 세션 — TIGER·ACE 이벤트 카드 화면 ]", width_cm=14, height_cm=6)
caption("그림 4.  경쟁사 채널 분석: TIGER·ACE·SOL 마케팅 이벤트 자동 감지 및 분류")

h3("③ 3단계 — 이중차분법(DiD)으로 순매수 효과 정량 측정")

body(
    "채널별 마케팅 감지 결과에 KRX 투자자별 순매수 데이터를 결합해, "
    "마케팅 활동이 실제 ETF 순매수에 미친 효과를 이중차분법(DiD, Difference-in-Differences)으로 정량 측정한다."
)

colored_box(
    "DiD 공식 (증권·개인 채널)\n"
    "  · KODEX 변화율  =  (이번 주 순매수 - 베이스라인) ÷ 절대평균(mabs)\n"
    "  · 비교군 변화율  =  유사 ETF 군 평균 변화율\n"
    "  · DiD 값  =  KODEX 변화율 - 비교군 변화율\n"
    "\n"
    "DiD 값 > 0  →  마케팅 기간 중 비교군 대비 KODEX가 더 많이 순매수됨",
    bg='EFF6FF', border='1F6FEB'
)

body(
    "은행 채널은 신호가 작고 변동성이 낮아 2단계 Z-score 방식을 추가 적용한다. "
    "1단계 DiD 값을 과거 16주 이력의 평균·표준편차로 표준화하여 통계적 이상치를 감지한다."
)

colored_box(
    "은행 2단계 Z-score 공식\n"
    "  Z = (이번 주 DiD - 16주 평균) ÷ (16주 표준편차 + 0.01)\n"
    "\n"
    "  Z ≥ 2.0  🟢  강한 이상 감지 — 은행 마케팅 거의 확실\n"
    "  Z ≥ 1.0  🟡  이상 감지 — 역추적 권고\n"
    "  |Z| < 1.0  ⚪  정상 변동 범위\n"
    "  Z < -1.0  🔴  경쟁사 우위 — 경쟁 마케팅 의심",
    bg='F0FFF4', border='3FB950'
)

img_placeholder("[ 증권사 세션 — DiD 계산 결과 화면 ]", width_cm=14, height_cm=6)
caption("그림 5.  증권사 채널 DiD 분석: 마케팅 이벤트별 금융투자 순매수 효과 정량화")

img_placeholder("[ 은행 세션 — Z-score DiD 결과 화면 ]", width_cm=14, height_cm=6)
caption("그림 6.  은행 채널 2단계 DiD 분석: Z-score 기반 마케팅 이상 감지")

h3("④ 4단계 — 주간 종합 리포트 자동 생성")

body(
    "4개 채널 분석 결과와 KRX 수급 데이터를 통합하여 Claude AI가 주간 종합 리포트를 자동 생성한다. "
    "리포트는 마케팅 담당자가 바로 활용할 수 있도록 아래 5개 섹션으로 구성된다."
)

for item in [
    "이번 주 핵심 요약 (3줄 이내 가장 중요한 시사점)",
    "채널별 마케팅 활동 (증권·은행·개인·경쟁사 감지 이벤트 정리)",
    "수급 시그널 (투자자별 순매수 흐름의 주목할 패턴)",
    "경쟁사 동향 (KODEX에 영향 줄 경쟁사 마케팅 현황)",
    "다음 주 액션 제안 (우선순위 순 3~5개 구체적 액션)",
]:
    bullet(item)

body("리포트는 생성 후 자동 저장되어 다음 접속 시 버튼 없이 바로 불러온다.")

img_placeholder("[ 주간 종합 리포트 — 생성된 리포트 화면 ]", width_cm=14, height_cm=7)
caption("그림 7.  주간 종합 리포트: 6개 채널 데이터 통합 분석 및 마케팅 액션 제안 자동 생성")

divider()

# ── 4. 운영 설계 ─────────────────────────────────────────────────────────────
h2("4.  주간 운영 설계 — 담당자 업무 흐름")

body("AI Agent 도입 후 마케팅 담당자의 주간 업무 흐름을 다음과 같이 정립할 수 있다.")

tbl2 = doc.add_table(rows=5, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

sched = [
    ('시점', '담당자 액션', 'Agent 역할'),
    ('금요일 장 마감 후', 'KRX 수집 버튼 클릭 (VPN 연결)', 'ETF 투자자별 순매수 데이터 저장'),
    ('금요일 저녁', '전체 채널 수집 버튼 클릭', '4개 세션 자동 수집 + LLM 분석 + 저장'),
    ('월요일 오전', '주간 리포트 확인', '채널별 이벤트 카드·DiD 결과·액션 제안 표시'),
    ('월요일 오후', '주간 마케팅 방향 결정', '경쟁사 이벤트·수급 시그널 기반 대응 기획'),
]

for ri, row_data in enumerate(sched):
    for ci, txt in enumerate(row_data):
        cell = tbl2.cell(ri, ci)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9.5)
        if ri == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1F6FEB')
            tcPr.append(shd)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

divider()

# ── 5. 기대 효과 ──────────────────────────────────────────────────────────────
h2("5.  기대 효과 및 향후 확장 방향")

h3("기대 효과")
bullet("마케팅 모니터링 시간 단축: 수동 채널 확인(주 5~8시간) → Agent 자동화(주 30분 이내)")
bullet("경쟁사 대응 속도 향상: TIGER·ACE 이벤트 감지 당일 확인 → KODEX 대응 기획 리드타임 단축")
bullet("마케팅 효과 정량화: DiD 기반 수치로 이벤트 ROI 측정 및 채널별 비교 가능")
bullet("데이터 기반 의사결정: 직관이 아닌 16주 이력 기반 통계적 이상 감지로 액션 우선순위 결정")

h3("향후 확장 방향")
bullet("자동 알림: 경쟁사 주요 이벤트 감지 시 Slack·이메일 즉시 알림")
bullet("KRX 완전 자동화: 장 마감 후 자동 수집으로 담당자 개입 최소화")
bullet("챗봇 연동: 담당자가 '이번 주 TIGER 이벤트 알려줘' 등 자연어로 Agent에 질의")
bullet("성과 추적 대시보드: 채널별 DiD 값 누적 트렌드 시각화로 장기 마케팅 효과 분석")

divider()

# ── 결론 ──────────────────────────────────────────────────────────────────────
h2("결론")

colored_box(
    "본 팀은 대고객 채널을 중심으로 증권·은행·경쟁사를 아우르는\n"
    "ETF 마케팅 모니터링 AI Agent를 직접 설계·개발하였다.\n"
    "\n"
    "매주 4개 채널 자동 수집 → Claude LLM 분석 → DiD 효과 측정 → 리포트 생성까지\n"
    "담당자가 버튼 2번만 누르면 주간 마케팅 인사이트가 자동으로 산출된다.\n"
    "\n"
    "이 시스템을 통해 KODEX는 경쟁사 마케팅에 더 빠르게 대응하고,\n"
    "데이터에 근거한 마케팅 의사결정 체계를 구축할 수 있다.",
    bg='EFF6FF', border='1F6FEB'
)

doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_footer.add_run("삼성자산운용 ETF 마케팅 모니터링 AI Agent · Powered by Claude · 2026")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x6E, 0x76, 0x81)

# ── 저장 ─────────────────────────────────────────────────────────────────────
out = "C:/Users/USER/Desktop/2.3.7_AI_Agent_구축_활용방안.docx"
doc.save(out)
print(f"저장 완료: {out}")
